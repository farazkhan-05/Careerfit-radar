from __future__ import annotations

from dataclasses import dataclass
from io import BytesIO
from pathlib import Path
from typing import Final

from backend.utils.hash_utils import sha256_text
from backend.utils.text_utils import normalize_text

PDF_CONTENT_TYPES: Final[set[str]] = {
    "application/pdf",
    "application/x-pdf",
}
DOCX_CONTENT_TYPES: Final[set[str]] = {
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
}
TEXT_CONTENT_TYPES: Final[set[str]] = {"text/plain"}
SUPPORTED_CONTENT_TYPES: Final[set[str]] = (
    PDF_CONTENT_TYPES | DOCX_CONTENT_TYPES | TEXT_CONTENT_TYPES
)


class ResumeParseError(ValueError):
    pass


@dataclass(frozen=True)
class ParsedResumeChunk:
    chunk_index: int
    content: str
    text_hash: str


@dataclass(frozen=True)
class ParsedResume:
    file_name: str
    content_type: str
    parsed_text: str
    text_hash: str
    chunks: tuple[ParsedResumeChunk, ...]


def parse_resume_bytes(file_name: str, content_type: str, payload: bytes) -> ParsedResume:
    if not file_name.strip():
        raise ResumeParseError("Resume file name is required.")
    if not payload:
        raise ResumeParseError("Resume file is empty.")

    normalized_content_type = content_type.lower().split(";", maxsplit=1)[0].strip()
    extracted_text = _extract_text(file_name, normalized_content_type, payload)
    parsed_text = normalize_text(extracted_text)
    if not parsed_text:
        raise ResumeParseError("Resume did not contain extractable text.")

    return ParsedResume(
        file_name=file_name,
        content_type=normalized_content_type,
        parsed_text=parsed_text,
        text_hash=sha256_text(parsed_text),
        chunks=chunk_resume_text(parsed_text),
    )


def chunk_resume_text(
    text: str,
    *,
    max_chars: int = 1800,
    overlap_chars: int = 180,
) -> tuple[ParsedResumeChunk, ...]:
    if max_chars < 200:
        raise ValueError("max_chars must be at least 200.")
    if overlap_chars < 0 or overlap_chars >= max_chars:
        raise ValueError("overlap_chars must be non-negative and smaller than max_chars.")

    normalized = normalize_text(text)
    if not normalized:
        return ()

    chunks: list[ParsedResumeChunk] = []
    seen_hashes: set[str] = set()
    start = 0
    while start < len(normalized):
        end = min(start + max_chars, len(normalized))
        if end < len(normalized):
            boundary = normalized.rfind(" ", start, end)
            if boundary > start + max_chars // 2:
                end = boundary

        chunk_text = normalized[start:end].strip()
        chunk_hash = sha256_text(chunk_text)
        if chunk_text and chunk_hash not in seen_hashes:
            seen_hashes.add(chunk_hash)
            chunks.append(
                ParsedResumeChunk(
                    chunk_index=len(chunks),
                    content=chunk_text,
                    text_hash=chunk_hash,
                )
            )

        if end >= len(normalized):
            break

        next_start = max(0, end - overlap_chars)
        if next_start <= start:
            next_start = end
        start = next_start

    return tuple(chunks)


def _extract_text(file_name: str, content_type: str, payload: bytes) -> str:
    suffix = Path(file_name).suffix.lower()
    if content_type in PDF_CONTENT_TYPES or suffix == ".pdf":
        return _extract_pdf_text(payload)
    if content_type in DOCX_CONTENT_TYPES or suffix == ".docx":
        return _extract_docx_text(payload)
    if content_type in TEXT_CONTENT_TYPES or suffix in {".txt", ".text"}:
        return payload.decode("utf-8-sig")
    raise ResumeParseError(f"Unsupported resume type: {content_type or suffix}.")


def _extract_pdf_text(payload: bytes) -> str:
    try:
        from pypdf import PdfReader

        reader = PdfReader(BytesIO(payload))
        return "\n".join(page.extract_text() or "" for page in reader.pages)
    except Exception as exc:  # pragma: no cover - concrete exception types vary.
        raise ResumeParseError("Could not parse PDF resume.") from exc


def _extract_docx_text(payload: bytes) -> str:
    try:
        from docx import Document

        document = Document(BytesIO(payload))
        paragraphs = [paragraph.text for paragraph in document.paragraphs]
        table_cells = [
            cell.text
            for table in document.tables
            for row in table.rows
            for cell in row.cells
        ]
        return "\n".join(paragraphs + table_cells)
    except Exception as exc:  # pragma: no cover - concrete exception types vary.
        raise ResumeParseError("Could not parse DOCX resume.") from exc
