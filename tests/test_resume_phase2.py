from __future__ import annotations

from io import BytesIO
from uuid import uuid4

import pytest

from backend.services.candidate_profile_service import CandidateProfileService
from backend.services.resume_parser import ResumeParseError, chunk_resume_text, parse_resume_bytes


class FakeProfileClient:
    def generate_profile(self, prompt: str) -> dict[str, object]:
        assert "Resume:" in prompt
        return {
            "target_roles": [
                {"value": "Backend Engineer", "evidence": "Backend Engineer"},
                {"value": "Product Manager", "evidence": "Product Manager"},
            ],
            "technical_skills": [
                {"value": "Python", "evidence": "Python"},
                {"value": "Kubernetes", "evidence": "Kubernetes"},
            ],
            "tools": [{"value": "PostgreSQL", "evidence": "PostgreSQL"}],
            "domains": [{"value": "APIs", "evidence": "REST APIs"}],
            "soft_skills": [{"value": "Mentoring", "evidence": "mentored interns"}],
            "experience_years": 4.0,
            "experience_evidence": "4 years",
            "projects": [
                {
                    "name": "CareerFit Radar",
                    "summary": "Built ranking services",
                    "technologies": ["Python", "PostgreSQL"],
                    "evidence": "CareerFit Radar",
                },
                {
                    "name": "Unsupported",
                    "summary": "Made up project",
                    "technologies": [],
                    "evidence": "not in resume",
                },
            ],
        }


def test_parse_text_resume_and_chunks_are_deterministic() -> None:
    payload = b"Backend Engineer\nPython  PostgreSQL\nBuilt REST APIs."

    parsed = parse_resume_bytes("resume.txt", "text/plain", payload)
    parsed_again = parse_resume_bytes("resume.txt", "text/plain", payload)

    assert parsed.parsed_text == "Backend Engineer Python PostgreSQL Built REST APIs."
    assert parsed.text_hash == parsed_again.text_hash
    assert parsed.chunks == parsed_again.chunks
    assert len(parsed.chunks) == 1


def test_chunking_removes_duplicate_chunks() -> None:
    chunks = chunk_resume_text(
        "Python APIs. " * 30,
        max_chars=200,
        overlap_chars=0,
    )

    assert len({chunk.text_hash for chunk in chunks}) == len(chunks)


def test_parse_docx_resume() -> None:
    from docx import Document

    buffer = BytesIO()
    document = Document()
    document.add_paragraph("Backend Engineer")
    document.add_paragraph("Python and PostgreSQL")
    document.save(buffer)

    parsed = parse_resume_bytes(
        "resume.docx",
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        buffer.getvalue(),
    )

    assert "Backend Engineer" in parsed.parsed_text
    assert "Python and PostgreSQL" in parsed.parsed_text


def test_invalid_resume_files_fail_clearly() -> None:
    with pytest.raises(ResumeParseError, match="Unsupported resume type"):
        parse_resume_bytes("resume.png", "image/png", b"not a resume")

    with pytest.raises(ResumeParseError, match="Could not parse PDF resume"):
        parse_resume_bytes("resume.pdf", "application/pdf", b"not a pdf")


def test_candidate_profile_keeps_only_resume_supported_claims() -> None:
    resume_text = (
        "Backend Engineer with 4 years of experience using Python and PostgreSQL. "
        "Built REST APIs for CareerFit Radar and mentored interns."
    )
    profile = CandidateProfileService(FakeProfileClient()).extract_profile(
        resume_id=uuid4(),
        resume_text=resume_text,
    )

    assert profile.target_roles == ["Backend Engineer"]
    assert profile.skills["technical"] == ["Python"]
    assert profile.skills["tools"] == ["PostgreSQL"]
    assert profile.skills["domains"] == ["APIs"]
    assert profile.skills["soft"] == ["Mentoring"]
    assert profile.experience_years == 4.0
    assert [project["name"] for project in profile.projects] == ["CareerFit Radar"]
